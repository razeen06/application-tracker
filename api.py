import io
import json
import os
import re
from datetime import datetime, date, timedelta, timezone

import docx
from flask import Blueprint, request, jsonify, g, current_app
from google import genai
from google.genai import types as genai_types
from sqlalchemy import func
from sqlalchemy.exc import IntegrityError

import crypto
import email_matching
import gmail_client
from models import db, Application, ApplicationStatus, AISummary, CompanyProfile, ProcessedEmail
from auth import token_required

api_routes = Blueprint("api_routes", __name__, url_prefix="/api")

GEMINI_MODEL = "gemini-3.1-flash-lite"
MAX_PAGE_TEXT_CHARS = 3000
MAX_EMAIL_BODY_CHARS = 2000
MAX_RESUME_TEXT_CHARS = 8000
GEMINI_TIMEOUT_MS = 15000
COMPETITIVENESS_CACHE_TTL_DAYS = 30
APPLICATION_DISCOVERY_LOOKBACK_DAYS = {7, 30, 60, 90}
APPLICATION_DISCOVERY_CANDIDATES_PER_PAGE = 25
APPLICATION_DISCOVERY_BATCH_SIZE = 25

# Gmail does the cheap first pass. Keep this deliberately recall-first: many
# employers never send a literal confirmation, but a later assessment or
# progression email still proves that an application was submitted. Gemini
# remains the strict second pass, so broad subject/sender signals can improve
# recall without turning newsletters and unfinished applications into rows.
APPLICATION_DISCOVERY_GMAIL_QUERY = (
    '{"thank you for applying" "thanks for applying" '
    '"thank you for your application" "application received" '
    '"received your application" "application confirmation" '
    '"application submitted" "your application has been submitted" '
    '"successfully applied" "application acknowledgement" '
    '"application acknowledgment" "your application has progressed" '
    '"application has progressed" "moving on to the next stage" '
    '"next stage of our recruitment process" '
    '"next step in our selection process" '
    'subject:application subject:assessment subject:candidate '
    'from:myworkday.com from:workday.com from:greenhouse.io from:lever.co '
    'from:smartrecruiters.com from:icims.com from:ashbyhq.com '
    'from:pageuppeople.com from:pushapply.com from:fusiongc.com.au}'
)

APPLICATION_DISCOVERY_TITLE_STOP_WORDS = {
    "a", "an", "at", "for", "job", "of", "position", "role", "the"
}
APPLICATION_DISCOVERY_TITLE_ALIASES = {
    "engineering": "engineer",
    "internship": "intern",
    "internships": "intern",
    "programme": "program",
    "programmes": "program",
}
APPLICATION_DISCOVERY_COMPANY_SUFFIXES = {
    "corp", "corporation", "inc", "limited", "llc", "ltd", "plc", "pty"
}

# Resume upload: only these two -- Gemini's document understanding accepts
# PDF bytes directly (verified with a real API call), but explicitly
# rejects DOCX ("400 Unsupported MIME type"), so DOCX text is extracted
# locally via python-docx first. No other formats are accepted since
# neither path above covers them.
RESUME_ALLOWED_EXTENSIONS = {"pdf", "docx"}
MAX_RESUME_FILE_SIZE_BYTES = 10 * 1024 * 1024

# Application Funnel: how many days an application can sit at "Applied" with
# no further movement before the funnel counts it as "Ghosted" rather than
# "Awaiting Response". This is a display-time bucket only (see
# _compute_application_funnel) -- never written back to Application.status,
# since a company can still respond after this threshold and an application
# shouldn't be permanently mislabeled just because it crossed a date line.
GHOSTED_THRESHOLD_DAYS = 21

SUGGESTION_CATEGORIES = {"Interview Offered", "Action Required", "Progress", "Rejected"}
VALID_CLASSIFICATIONS = SUGGESTION_CATEGORIES | {"Not Relevant"}

# Maps an ai_suggested_status category onto the real ApplicationStatus it
# becomes when accepted. "Interview Offered" collapses onto the existing
# INTERVIEW status rather than getting its own near-duplicate value; the
# other three map 1:1 (ApplicationStatus gained ACTION_REQUIRED/PROGRESS
# members specifically to support this).
SUGGESTION_STATUS_MAP = {
    "Interview Offered": ApplicationStatus.INTERVIEW,
    "Action Required": ApplicationStatus.ACTION_REQUIRED,
    "Progress": ApplicationStatus.PROGRESS,
    "Rejected": ApplicationStatus.REJECTED,
}

_genai_client = None
_genai_client_checked = False


def _get_genai_client():
    # Lazy + cached: constructing the client is cheap, but this also lets a
    # missing key fail per-request (a clear "Summary unavailable") instead
    # of at import time, which would take down the whole app.
    global _genai_client, _genai_client_checked

    if not _genai_client_checked:
        _genai_client_checked = True
        api_key = os.getenv("GEMINI_API_KEY")
        if api_key:
            _genai_client = genai.Client(api_key=api_key)

    return _genai_client


def _parse_date(value):
    if not value:
        return None
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except ValueError:
        return None


@api_routes.route("/applications", methods=["GET"])
@token_required
def list_applications():
    applications = (
        Application.query
        .filter_by(user_id=g.current_user.email)
        .order_by(Application.id.desc())
        .all()
    )
    return jsonify([application.to_dict() for application in applications])


@api_routes.route("/applications", methods=["POST"])
@token_required
def create_application():
    data = request.get_json(silent=True) or {}

    title = data.get("title")
    company = data.get("company")

    if not title or not company:
        return jsonify({"error": "title and company are required"}), 400

    status_value = data.get("status", ApplicationStatus.APPLIED.value)

    try:
        status = ApplicationStatus(status_value)
    except ValueError:
        return jsonify({"error": f"Invalid status: {status_value}"}), 400

    suitability_score = data.get("suitability_score")
    competitiveness_score = data.get("competitiveness_score")

    application = Application(
        user_id=g.current_user.email,
        title=title,
        company=company,
        url=data.get("url"),
        status=status,
        applied_date=_parse_date(data.get("applied_date")) or date.today(),
        notes=data.get("notes", ""),
        suitability_score=suitability_score,
        competitiveness_score=competitiveness_score,
    )
    application.set_flags(data.get("flags", []))

    if suitability_score is not None and competitiveness_score is not None:
        # This new application hasn't been committed yet, so it can't
        # possibly be in its own comparable bucket -- no exclude needed.
        historical = _compute_historical_response_rate(g.current_user.email, competitiveness_score)
        application.priority_label = _compute_priority_label(
            suitability_score, competitiveness_score, historical[0] if historical else None
        )

    db.session.add(application)
    db.session.commit()

    return jsonify(application.to_dict()), 201


def _clear_suggestion(application):
    application.ai_suggested_status = None
    application.ai_suggestion_source_email_id = None
    application.ai_suggestion_seen = False
    application.ai_suggestion_created_at = None


@api_routes.route("/applications/<int:application_id>", methods=["PUT"])
@token_required
def update_application(application_id):
    application = Application.query.filter_by(
        id=application_id, user_id=g.current_user.email
    ).first()

    if not application:
        return jsonify({"error": "Application not found"}), 404

    data = request.get_json(silent=True) or {}

    if "title" in data:
        application.title = data["title"]
    if "company" in data:
        application.company = data["company"]
    if "url" in data:
        application.url = data["url"]
    if "notes" in data:
        application.notes = data["notes"]
    if "flags" in data:
        application.set_flags(data["flags"])
    if "applied_date" in data:
        parsed = _parse_date(data["applied_date"])
        if parsed:
            application.applied_date = parsed

    # Scores can arrive after the application was already tracked (e.g. the
    # popup's two Gemini calls settle at different times) -- accepting them
    # here, not just at creation, and recomputing priority_label whenever
    # either one changes keeps the stored label from going stale.
    rescored = False
    if "suitability_score" in data:
        application.suitability_score = data["suitability_score"]
        rescored = True
    if "competitiveness_score" in data:
        application.competitiveness_score = data["competitiveness_score"]
        rescored = True

    if rescored:
        if application.suitability_score is not None and application.competitiveness_score is not None:
            historical = _compute_historical_response_rate(
                g.current_user.email, application.competitiveness_score,
                exclude_application_id=application.id,
            )
            application.priority_label = _compute_priority_label(
                application.suitability_score, application.competitiveness_score,
                historical[0] if historical else None,
            )
        else:
            application.priority_label = None

    if data.get("accept_suggestion"):
        # "Accept as-is": copy the AI's suggested category into the real
        # status field and close out the suggestion. Distinct from the
        # "status" branch below because the suggestion category string
        # ("Interview Offered") doesn't always match the ApplicationStatus
        # value it maps to ("Interview") -- see SUGGESTION_STATUS_MAP.
        if not application.ai_suggested_status:
            return jsonify({"error": "No active suggestion to accept"}), 400
        application.status = SUGGESTION_STATUS_MAP[application.ai_suggested_status]
        _clear_suggestion(application)
    elif "status" in data:
        try:
            application.status = ApplicationStatus(data["status"])
        except ValueError:
            return jsonify({"error": f"Invalid status: {data['status']}"}), 400
        # A manual status edit -- whether or not it matches the AI's
        # suggestion -- counts as the user having resolved it. Clearing
        # unconditionally is harmless when there was no active suggestion
        # (the fields are already None).
        _clear_suggestion(application)
    elif "ai_suggestion_seen" in data:
        # Banner click: acknowledges the suggestion without confirming or
        # discarding it -- the suggestion itself stays active and visible.
        application.ai_suggestion_seen = bool(data["ai_suggestion_seen"])

    db.session.commit()

    return jsonify(application.to_dict())


@api_routes.route("/applications/<int:application_id>", methods=["DELETE"])
@token_required
def delete_application(application_id):
    application = Application.query.filter_by(
        id=application_id, user_id=g.current_user.email
    ).first()

    if not application:
        return jsonify({"error": "Application not found"}), 404

    # Status-scan history may point at this row. Keep the processed-message
    # record (so the same email is not analysed again) but release its
    # foreign key before deleting the application.
    ProcessedEmail.query.filter_by(application_id=application.id).update(
        {"application_id": None}, synchronize_session=False
    )
    db.session.delete(application)
    db.session.commit()

    return "", 204


@api_routes.route("/applications", methods=["DELETE"])
@token_required
def erase_all_applications():
    # Settings' "Erase all applications" -- a full, unscoped wipe for the
    # signed-in account. No undo; the client is expected to have already
    # confirmed with the user before calling this.
    ProcessedEmail.query.filter_by(user_id=g.current_user.id).update(
        {"application_id": None}, synchronize_session=False
    )
    Application.query.filter_by(user_id=g.current_user.email).delete()
    db.session.commit()
    return "", 204


# Every application starts life as "Applied" (Application.status's default),
# so the funnel is a single-hop fan-out from that one source node to
# whichever bucket each application currently sits in -- there's no stored
# history of status transitions to chart a deeper multi-hop path through, so
# a one-hop "Applied -> current bucket" fan-out is the honest shape rather
# than implying a specific journey we can't actually observe. Ordered
# roughly best-outcome-first; purely cosmetic (fixes rendering order in the
# frontend), doesn't affect the counts.
FUNNEL_STAGE_ORDER = [
    "Offer", "Interview", "Progress", "Action Required",
    "Awaiting Response", "Rejected", "Ghosted",
]


def _compute_application_funnel(user_email):
    # Pure SQL/ORM aggregation -- no AI involved, this is just counting.
    ghosted_cutoff = date.today() - timedelta(days=GHOSTED_THRESHOLD_DAYS)

    non_applied_counts = dict(
        db.session.query(Application.status, func.count(Application.id))
        .filter(Application.user_id == user_email, Application.status != ApplicationStatus.APPLIED)
        .group_by(Application.status)
        .all()
    )

    # "Ghosted" is a derived/virtual bucket, computed fresh on every call --
    # never stored on the row -- specifically because a company can still
    # respond after crossing this date line, so a stored label would risk
    # staying wrong after the fact.
    ghosted_count = Application.query.filter(
        Application.user_id == user_email,
        Application.status == ApplicationStatus.APPLIED,
        Application.applied_date.isnot(None),
        Application.applied_date <= ghosted_cutoff,
    ).count()

    still_applied_count = Application.query.filter(
        Application.user_id == user_email,
        Application.status == ApplicationStatus.APPLIED,
    ).count()
    # Applications with no applied_date at all can't have their age judged,
    # so they fall into "Awaiting Response" alongside ones too recent to be
    # ghosted, rather than being silently dropped from the funnel.
    awaiting_count = still_applied_count - ghosted_count

    stage_counts = {
        "Interview": non_applied_counts.get(ApplicationStatus.INTERVIEW, 0),
        "Action Required": non_applied_counts.get(ApplicationStatus.ACTION_REQUIRED, 0),
        "Progress": non_applied_counts.get(ApplicationStatus.PROGRESS, 0),
        "Offer": non_applied_counts.get(ApplicationStatus.OFFER, 0),
        "Rejected": non_applied_counts.get(ApplicationStatus.REJECTED, 0),
        "Ghosted": ghosted_count,
        "Awaiting Response": awaiting_count,
    }

    flows = [
        {"from_stage": "Applied", "to_stage": stage, "count": stage_counts[stage]}
        for stage in FUNNEL_STAGE_ORDER
        if stage_counts[stage] > 0
    ]

    total = sum(stage_counts.values())
    responded = (
        stage_counts["Interview"] + stage_counts["Action Required"]
        + stage_counts["Progress"] + stage_counts["Offer"] + stage_counts["Rejected"]
    )

    return {
        "total": total,
        "flows": flows,
        "stage_counts": stage_counts,
        "summary": {
            "total": total,
            "responded": responded,
            "ghosted": stage_counts["Ghosted"],
            "awaiting_response": stage_counts["Awaiting Response"],
        },
    }


@api_routes.route("/application-funnel", methods=["GET"])
@token_required
def get_application_funnel():
    return jsonify(_compute_application_funnel(g.current_user.email))


@api_routes.route("/background", methods=["GET"])
@token_required
def get_background():
    return jsonify({"background_text": g.current_user.background_text or ""})


@api_routes.route("/background", methods=["PUT"])
@token_required
def update_background():
    data = request.get_json(silent=True) or {}
    # Explicit key check (not just .get(..., "")) so omitting the key
    # entirely is a no-op rather than accidentally clearing it.
    if "background_text" not in data:
        return jsonify({"error": "background_text is required"}), 400

    g.current_user.background_text = (data["background_text"] or "").strip() or None
    db.session.commit()

    return jsonify({"background_text": g.current_user.background_text or ""})


RESUME_EXTRACTION_INSTRUCTION = (
    "You are extracting structured data from a resume. Read the attached "
    "resume content and respond with a JSON object containing exactly four "
    "keys:\n\n"
    "\"skills\": a JSON array of strings -- technical and professional "
    "skills explicitly listed or clearly demonstrated in the resume.\n\n"
    "\"education\": a JSON array of objects, one per degree/program, each "
    "with keys \"degree\", \"field\", and \"institution\" (use an empty "
    "string for any that aren't stated).\n\n"
    "\"work_experience\": a JSON array of objects, one per role, each with "
    "keys \"role\", \"company\", \"duration\", and \"description\" (a brief "
    "1-2 sentence summary of what they actually did in that role).\n\n"
    "\"interests\": a JSON array of strings -- explicitly stated interests, "
    "career goals, or the type of role they say they're looking for. Use an "
    "empty array if none are stated -- do not guess.\n\n"
    "Only include information that is actually present in the resume -- do "
    "not invent or infer details that aren't there."
)


def _extract_docx_text(file_bytes):
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # Some resume templates put work experience/skills in a table-based
    # layout rather than plain paragraphs -- pull those cells in too.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _clean_resume_entry(entry, keys):
    if not isinstance(entry, dict):
        return None
    cleaned = {}
    for key in keys:
        value = entry.get(key, "")
        cleaned[key] = value if isinstance(value, str) else ""
    return cleaned


def _parse_resume_extraction_response(raw_text):
    cleaned = raw_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        parsed = json.loads(cleaned)
    except (json.JSONDecodeError, TypeError):
        return None

    if not isinstance(parsed, dict):
        return None

    skills = parsed.get("skills")
    education = parsed.get("education")
    work_experience = parsed.get("work_experience")
    interests = parsed.get("interests")

    if not isinstance(skills, list) or not all(isinstance(s, str) for s in skills):
        return None
    if not isinstance(interests, list) or not all(isinstance(s, str) for s in interests):
        return None
    if not isinstance(education, list) or not isinstance(work_experience, list):
        return None

    cleaned_education = []
    for entry in education:
        cleaned_entry = _clean_resume_entry(entry, ("degree", "field", "institution"))
        if cleaned_entry is None:
            return None
        cleaned_education.append(cleaned_entry)

    cleaned_experience = []
    for entry in work_experience:
        cleaned_entry = _clean_resume_entry(entry, ("role", "company", "duration", "description"))
        if cleaned_entry is None:
            return None
        cleaned_experience.append(cleaned_entry)

    return {
        "skills": skills,
        "education": cleaned_education,
        "work_experience": cleaned_experience,
        "interests": interests,
    }


@api_routes.route("/resume", methods=["GET"])
@token_required
def get_resume():
    return jsonify({"resume_structured": g.current_user.get_resume_structured()})


@api_routes.route("/resume", methods=["POST"])
@token_required
def upload_resume():
    if "resume" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["resume"]
    filename = file.filename or ""
    if "." not in filename:
        return jsonify({"error": "Could not determine file type -- upload a PDF or DOCX"}), 400

    ext = filename.rsplit(".", 1)[1].lower()
    if ext not in RESUME_ALLOWED_EXTENSIONS:
        return jsonify({"error": "Only PDF and DOCX files are supported"}), 400

    file_bytes = file.read()
    if not file_bytes:
        return jsonify({"error": "The uploaded file is empty"}), 400
    if len(file_bytes) > MAX_RESUME_FILE_SIZE_BYTES:
        return jsonify({"error": "File is too large (max 10MB)"}), 400

    client = _get_genai_client()
    if client is None:
        current_app.logger.warning("Resume parsing skipped: GEMINI_API_KEY not configured")
        return jsonify({"error": "Resume parsing unavailable"}), 503

    # Processed entirely in memory and never written to disk or stored --
    # only the parsed structured result below is persisted (User.resume_
    # structured). Re-uploading is cheap for the user if parsing fails, so
    # there's no "allow re-parse without re-upload" case strong enough to
    # justify retaining a PII-bearing raw file, even temporarily.
    if ext == "pdf":
        contents = [
            genai_types.Part.from_bytes(data=file_bytes, mime_type="application/pdf"),
            RESUME_EXTRACTION_INSTRUCTION,
        ]
    else:
        try:
            resume_text = _extract_docx_text(file_bytes)
        except Exception as e:
            current_app.logger.warning(f"DOCX text extraction failed: {e}")
            return jsonify({"error": "Couldn't read that DOCX file -- it may be corrupted"}), 400
        if not resume_text.strip():
            return jsonify({"error": "Couldn't find any text in that DOCX file"}), 400
        contents = f"{RESUME_EXTRACTION_INSTRUCTION}\n\nResume text:\n{resume_text[:MAX_RESUME_TEXT_CHARS]}"

    try:
        result = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=contents,
            config=genai_types.GenerateContentConfig(
                http_options=genai_types.HttpOptions(timeout=GEMINI_TIMEOUT_MS),
                response_mime_type="application/json",
            ),
        )
        raw_text = (result.text or "").strip()
    except Exception as e:
        current_app.logger.warning(f"Gemini resume extraction call failed: {e}")
        return jsonify({"error": "Couldn't parse that resume -- try again"}), 503

    parsed = _parse_resume_extraction_response(raw_text)
    if parsed is None:
        current_app.logger.warning("Gemini resume extraction response wasn't valid structured JSON")
        return jsonify({"error": "Couldn't parse that resume -- try again, or use a different file"}), 503

    g.current_user.set_resume_structured(parsed)
    db.session.commit()

    return jsonify({"resume_structured": parsed})


@api_routes.route("/scan-history/reset", methods=["POST"])
@token_required
def reset_scan_history():
    # Settings' "Reset AI scan history" -- forgets which emails have already
    # been looked at (and the scan watermark), so the next scan re-evaluates
    # the account's whole applicable email history from scratch. Doesn't
    # touch any tracked application or its status.
    user = g.current_user
    ProcessedEmail.query.filter_by(user_id=user.id).delete()
    user.last_email_scan_at = None
    db.session.commit()
    return "", 204


def _build_summary_prompt(page_text):
    truncated = page_text[:MAX_PAGE_TEXT_CHARS]
    return (
        "Analyze this job posting and respond with a JSON object containing "
        "exactly two keys: \"summary\" and \"flags\".\n\n"
        "\"summary\": a 3-bullet plain-text summary, as a single string "
        "(use \\n between bullets -- no markdown, no asterisks, no headers, "
        "since this renders in a small browser extension popup). Cover "
        "exactly these three points in this order:\n"
        "1) The role and key responsibilities.\n"
        "2) Pay or compensation, if mentioned (write \"Not mentioned\" if it isn't).\n"
        "3) Eligibility requirements such as year level, WAM/GPA cutoff, or "
        "visa status, if mentioned (write \"None mentioned\" if there aren't any).\n\n"
        "\"flags\": a JSON array of zero or more of these exact label "
        "strings -- include a label only if the posting genuinely matches it:\n"
        "- \"Unpaid internship\": only if the role is unpaid AND no real "
        "salary/pay figure (e.g. \"$20/hour\", \"$50,000/year\") is stated "
        "anywhere in the posting. If a specific pay amount is mentioned "
        "anywhere, do NOT include this label even if the word \"unpaid\" "
        "appears elsewhere in the text (e.g. in an unrelated leave/policy line).\n"
        "- \"WAM/GPA cutoff mentioned\": if a minimum WAM or GPA requirement is stated.\n"
        "- \"'Penultimate year' requirement\": if the posting requires penultimate-year status.\n"
        "- \"'Final year' requirement\": if the posting requires final-year status.\n"
        "- \"Citizenship/visa restriction\": if eligibility is restricted by "
        "citizenship, permanent residency, or visa status.\n"
        "Use an empty array for \"flags\" if none apply.\n\n"
        f"Job posting text:\n{truncated}"
    )


def _parse_summary_response(raw_text, fallback_flags):
    # Gemini is asked for plain JSON, but LLMs sometimes wrap it in a
    # markdown code fence anyway despite instructions -- strip that first.
    cleaned = raw_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        parsed = json.loads(cleaned)
    except (json.JSONDecodeError, TypeError):
        parsed = None

    if (
        isinstance(parsed, dict)
        and isinstance(parsed.get("summary"), str)
        and parsed["summary"].strip()
        and isinstance(parsed.get("flags"), list)
        and all(isinstance(flag, str) for flag in parsed["flags"])
    ):
        return parsed["summary"].strip(), parsed["flags"]

    # Malformed/unexpected shape -- fall back to the old plain-text-summary
    # behavior, and to the client's own regex-derived flags rather than
    # silently dropping flag detection entirely.
    current_app.logger.warning(
        "Gemini summarize response wasn't valid {summary, flags} JSON -- "
        "falling back to plain text + client-supplied flags"
    )
    return raw_text, fallback_flags


@api_routes.route("/summarize", methods=["POST"])
@token_required
def summarize():
    data = request.get_json(silent=True) or {}

    url = data.get("url")
    page_text = data.get("page_text")

    if not url or not page_text:
        return jsonify({"error": "url and page_text are required"}), 400

    existing = AISummary.query.filter_by(url=url).first()
    if existing:
        response = existing.to_dict()
        response["cached"] = True
        return jsonify(response)

    client = _get_genai_client()
    if client is None:
        current_app.logger.warning("Gemini summarize skipped: GEMINI_API_KEY not configured")
        return jsonify({"error": "Summary unavailable"}), 503

    try:
        result = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=_build_summary_prompt(page_text),
            config=genai_types.GenerateContentConfig(
                http_options=genai_types.HttpOptions(timeout=GEMINI_TIMEOUT_MS),
                # First line of defense for getting clean JSON back -- the
                # parse/validate/fallback below is the second, since even
                # JSON mode isn't a hard guarantee of the exact {summary,
                # flags} shape we asked for.
                response_mime_type="application/json",
            ),
        )
        raw_text = (result.text or "").strip()
    except Exception as e:
        current_app.logger.warning(f"Gemini summarize call failed for {url}: {e}")
        return jsonify({"error": "Summary unavailable"}), 503

    if not raw_text:
        current_app.logger.warning(f"Gemini summarize returned empty text for {url}")
        return jsonify({"error": "Summary unavailable"}), 503

    # AI-first flag detection: prefer Gemini's own analysis (which can
    # reason about context, e.g. an incidental "unpaid leave" mention
    # alongside a real salary), falling back to the extension's regex-based
    # flags (passed in the request) only if the response didn't parse.
    client_flags = data.get("flags", [])
    summary_text, flags = _parse_summary_response(raw_text, client_flags)

    summary = AISummary(url=url, summary_text=summary_text)
    summary.set_flags_snapshot(flags)
    db.session.add(summary)

    try:
        db.session.commit()
    except IntegrityError:
        # A concurrent request for the same URL won the race and already
        # inserted a row (url is unique) -- the Gemini call above already
        # happened and can't be un-billed, but at least don't error out or
        # store a duplicate; return the row that actually won.
        db.session.rollback()
        summary = AISummary.query.filter_by(url=url).first()

    response = summary.to_dict()
    response["cached"] = False
    return jsonify(response), 201


def _format_resume_for_prompt(resume_structured):
    lines = []

    skills = resume_structured.get("skills") or []
    if skills:
        lines.append("Skills: " + ", ".join(skills))

    education = resume_structured.get("education") or []
    if education:
        lines.append("Education:")
        for entry in education:
            degree_parts = [p for p in (entry.get("degree"), entry.get("field")) if p]
            degree_line = " in ".join(degree_parts) if degree_parts else "Degree"
            institution = entry.get("institution")
            lines.append(f"- {degree_line}" + (f", {institution}" if institution else ""))

    work_experience = resume_structured.get("work_experience") or []
    if work_experience:
        lines.append("Work experience:")
        for entry in work_experience:
            header = entry.get("role") or "Role"
            if entry.get("company"):
                header += f" at {entry['company']}"
            if entry.get("duration"):
                header += f" ({entry['duration']})"
            lines.append(f"- {header}")
            if entry.get("description"):
                lines.append(f"  {entry['description']}")

    interests = resume_structured.get("interests") or []
    if interests:
        lines.append("Stated interests/career goals: " + ", ".join(interests))

    return "\n".join(lines)


def _build_candidate_profile_text(resume_structured, background_text):
    # resume_structured is the primary input once a resume's been uploaded;
    # background_text becomes supplementary -- anything the resume doesn't
    # capture, like a role preference typed in separately (see models.py's
    # User.background_text). Falls back to background_text alone if no
    # resume exists, and combines both when both do -- the prompt itself
    # just needs one "candidate's background" block either way.
    parts = []
    if resume_structured:
        parts.append("Resume:\n" + _format_resume_for_prompt(resume_structured))
        if background_text:
            parts.append(f"Additional notes from the candidate:\n{background_text}")
    elif background_text:
        parts.append(f"Candidate background:\n{background_text}")
    return "\n\n".join(parts)


def _build_suitability_prompt(candidate_profile_text, page_text):
    truncated = page_text[:MAX_PAGE_TEXT_CHARS]
    return (
        "You are helping a job seeker judge how well they personally suit a "
        "specific job posting, based on their own background.\n\n"
        f"Candidate's background:\n{candidate_profile_text}\n\n"
        f"Job posting:\n{truncated}\n\n"
        "Respond with a JSON object containing exactly two keys:\n"
        "\"score\": a number from 0 to 10 rating how well this candidate's "
        "background suits this specific role (0 = no meaningful overlap, "
        "10 = an excellent match). Judge honestly based on actual overlap "
        "between their background and the role's actual requirements -- a "
        "genuine mismatch should score low (e.g. 1-2), not a safe middle "
        "value, and a strong match should score high (e.g. 8-10).\n"
        "\"rationale\": one sentence explaining the score.\n"
    )


def _parse_suitability_response(raw_text):
    cleaned = raw_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        parsed = json.loads(cleaned)
    except (json.JSONDecodeError, TypeError):
        return None

    if not isinstance(parsed, dict):
        return None

    score = parsed.get("score")
    rationale = parsed.get("rationale")

    if not isinstance(score, (int, float)) or isinstance(score, bool):
        return None
    if not isinstance(rationale, str) or not rationale.strip():
        return None

    return {"score": max(0.0, min(10.0, float(score))), "rationale": rationale.strip()}


@api_routes.route("/score-suitability", methods=["POST"])
@token_required
def score_suitability():
    data = request.get_json(silent=True) or {}
    page_text = data.get("page_text")

    if not page_text:
        return jsonify({"error": "page_text is required"}), 400

    resume_structured = g.current_user.get_resume_structured()
    background_text = g.current_user.background_text
    if not resume_structured and not background_text:
        # Not a hard error -- this is an expected, common state (most users
        # haven't filled in Settings yet). No score to guess at without one.
        return jsonify({
            "suitability_score": None,
            "suitability_rationale": None,
            "message": "Add your background or upload a resume in Settings for a suitability score",
        })

    candidate_profile_text = _build_candidate_profile_text(resume_structured, background_text)

    client = _get_genai_client()
    if client is None:
        current_app.logger.warning("Gemini suitability scoring skipped: GEMINI_API_KEY not configured")
        return jsonify({"error": "Suitability score unavailable"}), 503

    try:
        result = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=_build_suitability_prompt(candidate_profile_text, page_text),
            config=genai_types.GenerateContentConfig(
                http_options=genai_types.HttpOptions(timeout=GEMINI_TIMEOUT_MS),
                response_mime_type="application/json",
            ),
        )
        raw_text = (result.text or "").strip()
    except Exception as e:
        current_app.logger.warning(f"Gemini suitability scoring call failed: {e}")
        return jsonify({"error": "Suitability score unavailable"}), 503

    parsed = _parse_suitability_response(raw_text)
    if parsed is None:
        current_app.logger.warning("Gemini suitability response wasn't valid {score, rationale} JSON")
        return jsonify({"error": "Suitability score unavailable"}), 503

    return jsonify({"suitability_score": parsed["score"], "suitability_rationale": parsed["rationale"]})


def _build_competitiveness_prompt(company_name):
    return (
        f"Assess how competitive it is for a student or early-career "
        f"candidate to get an internship or entry-level job offer at "
        f"\"{company_name}\", on a scale of 0 to 10 -- where 0 means almost "
        "any reasonable applicant gets an offer, and 10 means it's among the "
        "most selective employers to get into (on par with top-tier tech "
        "companies, elite finance/consulting firms, etc.).\n\n"
        "Respond with a JSON object containing exactly two keys:\n"
        "\"score\": a number from 0 to 10.\n"
        "\"rationale\": one sentence explaining the score.\n"
    )


def _parse_competitiveness_response(raw_text):
    # Same {score, rationale} shape as suitability -- reuse the same parsing
    # rules (markdown-fence stripping, 0-10 clamping) rather than duplicating
    # them under a different name.
    return _parse_suitability_response(raw_text)


def _fetch_competitiveness_from_gemini(client, company_name):
    # Try a grounded (real Google Search) call first, so the score reflects
    # actual current information about the company rather than the model's
    # training data alone. Whether this is usable depends on the API key's
    # tier -- some keys hit a quota error specific to the search tool even
    # though plain calls work fine, so this always falls back to a plain
    # prompt rather than surfacing that as a hard failure. Either way, the
    # `grounded` flag reflects what actually happened, not what was
    # attempted, so an ungrounded guess is never presented as verified
    # research (see CompanyProfile.grounded).
    prompt = _build_competitiveness_prompt(company_name)

    try:
        result = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt + "\nRespond with ONLY the JSON object, no other text.",
            config=genai_types.GenerateContentConfig(
                http_options=genai_types.HttpOptions(timeout=GEMINI_TIMEOUT_MS),
                tools=[genai_types.Tool(google_search=genai_types.GoogleSearch())],
            ),
        )
        raw_text = (result.text or "").strip()
        parsed = _parse_competitiveness_response(raw_text)
        if parsed is not None:
            grounding_metadata = (
                result.candidates[0].grounding_metadata if result.candidates else None
            )
            actually_grounded = bool(
                grounding_metadata and grounding_metadata.grounding_chunks
            )
            return parsed["score"], parsed["rationale"], actually_grounded
        current_app.logger.warning(
            f"Grounded competitiveness response for {company_name} wasn't valid JSON -- falling back"
        )
    except Exception as e:
        current_app.logger.warning(
            f"Grounded competitiveness call failed for {company_name}, falling back to plain: {e}"
        )

    try:
        result = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt,
            config=genai_types.GenerateContentConfig(
                http_options=genai_types.HttpOptions(timeout=GEMINI_TIMEOUT_MS),
                response_mime_type="application/json",
            ),
        )
        raw_text = (result.text or "").strip()
    except Exception as e:
        current_app.logger.warning(f"Plain competitiveness call failed for {company_name}: {e}")
        return None

    parsed = _parse_competitiveness_response(raw_text)
    if parsed is None:
        current_app.logger.warning(
            f"Plain competitiveness response for {company_name} wasn't valid {{score, rationale}} JSON"
        )
        return None

    return parsed["score"], parsed["rationale"], False


@api_routes.route("/score-competitiveness", methods=["POST"])
@token_required
def score_competitiveness():
    data = request.get_json(silent=True) or {}
    company_name = (data.get("company_name") or "").strip()

    if not company_name:
        return jsonify({"error": "company_name is required"}), 400

    stale_cutoff = datetime.utcnow() - timedelta(days=COMPETITIVENESS_CACHE_TTL_DAYS)
    existing = CompanyProfile.query.filter_by(company_name=company_name).first()

    if existing and existing.fetched_at >= stale_cutoff:
        response = existing.to_dict()
        response["cached"] = True
        return jsonify(response)

    client = _get_genai_client()
    if client is None:
        current_app.logger.warning("Gemini competitiveness scoring skipped: GEMINI_API_KEY not configured")
        return jsonify({"error": "Competitiveness score unavailable"}), 503

    result = _fetch_competitiveness_from_gemini(client, company_name)
    if result is None:
        return jsonify({"error": "Competitiveness score unavailable"}), 503

    score, rationale, grounded = result

    if existing:
        existing.competitiveness_score = score
        existing.rationale = rationale
        existing.grounded = grounded
        existing.fetched_at = datetime.utcnow()
    else:
        existing = CompanyProfile(
            company_name=company_name,
            competitiveness_score=score,
            rationale=rationale,
            grounded=grounded,
        )
        db.session.add(existing)

    try:
        db.session.commit()
    except IntegrityError:
        # A concurrent request for the same brand-new company won the race
        # and already inserted a row (company_name is unique) -- the Gemini
        # call above already happened and can't be un-billed, but at least
        # don't error out; return the row that actually won.
        db.session.rollback()
        existing = CompanyProfile.query.filter_by(company_name=company_name).first()

    response = existing.to_dict()
    response["cached"] = False
    return jsonify(response), 201


# Historical-performance signal for priority scoring -- deliberately NOT an
# AI call. Computed directly from the user's own past Application rows: of
# the applications they've tracked at a comparable competitiveness level
# (+/- COMPETITIVENESS_BUCKET_RADIUS), what fraction ever moved past
# "Applied"? This is a response rate (did the company engage at all), not an
# advancement rate -- Interview, Action Required, Progress, Offer, and
# Rejected all count as "responded", since all five mean the company did
# something with the application; only "Applied" with no further movement
# counts as no response yet.
COMPETITIVENESS_BUCKET_RADIUS = 2.0
MIN_HISTORICAL_BUCKET_SIZE = 3


def _compute_historical_response_rate(user_email, competitiveness_score, exclude_application_id=None):
    # Returns (response_rate: float in [0, 1], sample_size: int), or None if
    # there isn't enough comparable history yet to say anything meaningful.
    if competitiveness_score is None:
        return None

    query = Application.query.filter(
        Application.user_id == user_email,
        Application.competitiveness_score.isnot(None),
        Application.competitiveness_score >= competitiveness_score - COMPETITIVENESS_BUCKET_RADIUS,
        Application.competitiveness_score <= competitiveness_score + COMPETITIVENESS_BUCKET_RADIUS,
    )
    if exclude_application_id is not None:
        query = query.filter(Application.id != exclude_application_id)

    comparable = query.all()
    if len(comparable) < MIN_HISTORICAL_BUCKET_SIZE:
        return None

    responded = sum(1 for a in comparable if a.status != ApplicationStatus.APPLIED)
    return responded / len(comparable), len(comparable)


# Application Priority label: a single deterministic, weighted combination
# of suitability (AI, how well this candidate fits this role),
# competitiveness (AI/cached, how hard this company is to get into), and the
# user's own historical response rate at a comparable competitiveness level
# (real, from _compute_historical_response_rate -- folded in only when
# there's enough of it). No AI call decides the label itself -- fixed
# thresholds on a fixed formula, confirmed with the user before finalizing
# the wording.
#
# Ordered lowest to highest priority; checked top-down so the first
# threshold a score clears wins.
PRIORITY_LABEL_THRESHOLDS = [
    (8.0, "Top Priority"),
    (6.0, "Strong Match"),
    (4.0, "Worth Applying"),
    (0.0, "Low Priority"),
]


def _label_for_combined_score(combined_score):
    for threshold, label in PRIORITY_LABEL_THRESHOLDS:
        if combined_score >= threshold:
            return label
    return PRIORITY_LABEL_THRESHOLDS[-1][1]


def _compute_priority_label(suitability_score, competitiveness_score, historical_response_rate):
    # Both AI scores are required inputs -- without a suitability score
    # (background_text unset) or a competitiveness score (not yet fetched),
    # there isn't enough to combine, so no label rather than a guess built
    # on a missing input.
    if suitability_score is None or competitiveness_score is None:
        return None

    if historical_response_rate is not None:
        combined = (
            0.4 * suitability_score
            + 0.3 * (10 - competitiveness_score)
            + 0.3 * (historical_response_rate * 10)
        )
    else:
        # Cold start: no comparable history yet, so the two AI scores split
        # the full weight rather than historical_response_rate defaulting to
        # a fabricated value.
        combined = 0.5 * suitability_score + 0.5 * (10 - competitiveness_score)

    return _label_for_combined_score(combined)


@api_routes.route("/compute-priority", methods=["POST"])
@token_required
def compute_priority():
    # Pure computation, no persistence -- lets the extension popup show a
    # priority label as soon as it has both scores in hand, before the user
    # has necessarily tracked the application yet (see popup.js).
    data = request.get_json(silent=True) or {}
    suitability_score = data.get("suitability_score")
    competitiveness_score = data.get("competitiveness_score")

    historical = None
    if competitiveness_score is not None:
        historical = _compute_historical_response_rate(g.current_user.email, competitiveness_score)

    historical_rate = historical[0] if historical else None
    label = _compute_priority_label(suitability_score, competitiveness_score, historical_rate)

    return jsonify({
        "priority_label": label,
        "historical_response_rate": historical_rate,
        "historical_sample_size": historical[1] if historical else 0,
    })


def _build_classification_prompt(subject, body):
    truncated = body[:MAX_EMAIL_BODY_CHARS]
    return (
        "You are helping a job seeker track their internship/job applications. "
        "Read this email and classify it into EXACTLY ONE of these five "
        "categories, based on what it means for the specific application it "
        "was matched to (the match was done by a separate, loose subject-line "
        "heuristic, so double check the body actually is about this "
        "application before picking anything other than \"Not Relevant\"):\n\n"
        "- \"Interview Offered\": invites the candidate to interview, schedule "
        "a call, or take an assessment/OA as the next step.\n"
        "- \"Action Required\": asks the candidate to do something (complete a "
        "form, provide documents, confirm details) that isn't an interview invite.\n"
        "- \"Progress\": a positive status update (e.g. application received/"
        "under review, moved to the next round) with no concrete next action yet.\n"
        "- \"Rejected\": states the candidate was not selected, or the role was "
        "filled/closed.\n"
        "- \"Not Relevant\": not actually about this specific job application "
        "(a newsletter, an unrelated email that happens to mention the company "
        "or role in passing, a promotional email, etc.).\n\n"
        "Respond with a JSON object containing exactly one key, "
        "\"classification\", whose value is one of the five exact strings above.\n\n"
        f"Email subject: {subject}\n\n"
        f"Email body:\n{truncated}"
    )


def _parse_classification_response(raw_text):
    cleaned = raw_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        parsed = json.loads(cleaned)
    except (json.JSONDecodeError, TypeError):
        return None

    if isinstance(parsed, dict) and parsed.get("classification") in VALID_CLASSIFICATIONS:
        return parsed["classification"]

    return None


def _build_application_discovery_prompt(messages):
    email_blocks = []
    for message in messages:
        email_blocks.append(
            "\n".join([
                f"Message ID: {message['id']}",
                f"Subject: {message['subject']}",
                f"Sender: {message['sender']}",
                f"Received date: {message['received_date']}",
                f"Body: {message['body'][:MAX_EMAIL_BODY_CHARS]}",
            ])
        )

    return (
        "You are identifying historical job applications from application-"
        "related emails. For each email below, decide whether it clearly confirms "
        "that the recipient submitted an application for a job, internship, "
        "graduate program, or similar role. A later-stage email is valid "
        "evidence when it explicitly says the recipient's application has "
        "progressed, or clearly invites them to an assessment or next stage "
        "for a named employer and role/program. It does not need to be the "
        "original submission confirmation.\n\n"
        "Do not count newsletters, job alerts, recruiter outreach unrelated "
        "to an existing application, generic account creation or candidate "
        "verification, expressions of interest, or reminders saying an "
        "application is still in progress, incomplete, or has not yet been "
        "submitted. An invitation to finish or submit an application is not "
        "proof that it was submitted.\n\n"
        "Respond with a JSON object containing exactly one key, "
        '"applications". Its value must be an array containing only confirmed '
        "applications. Each item must contain exactly these string keys:\n"
        '- "message_id": copy the supplied Message ID exactly.\n'
        '- "company": the employer or organisation.\n'
        '- "title": the role/program title.\n'
        '- "applied_date": YYYY-MM-DD, using the explicit submission date if '
        "stated, otherwise the supplied received date.\n\n"
        "If several emails refer to the same application, return only the "
        "clearest one, preferring a direct confirmation over a reminder. Be "
        "conservative: if an email is ambiguous, leave it out. Never invent a "
        "company or role.\n\n"
        + "\n\n--- EMAIL ---\n".join(email_blocks)
    )


def _parse_application_discovery_response(raw_text, allowed_message_ids):
    cleaned = raw_text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        parsed = json.loads(cleaned)
    except (json.JSONDecodeError, TypeError):
        return None

    applications = parsed.get("applications") if isinstance(parsed, dict) else None
    if not isinstance(applications, list):
        return None

    validated = []
    seen_message_ids = set()
    for item in applications:
        if not isinstance(item, dict):
            return None
        message_id = item.get("message_id")
        company = item.get("company")
        title = item.get("title")
        applied_date = item.get("applied_date")
        if (
            not isinstance(message_id, str)
            or message_id not in allowed_message_ids
            or message_id in seen_message_ids
            or not isinstance(company, str)
            or not company.strip()
            or not isinstance(title, str)
            or not title.strip()
            or not isinstance(applied_date, str)
        ):
            continue
        seen_message_ids.add(message_id)
        validated.append({
            "message_id": message_id,
            "company": company.strip()[:200],
            "title": title.strip()[:200],
            "applied_date": applied_date,
        })
    return validated


def _discovery_normalize(value):
    return re.sub(r"[^a-z0-9]+", " ", (value or "").lower()).strip()


def _discovery_company_key(value):
    tokens = _discovery_normalize(value).split()
    if tokens and tokens[0] == "the":
        tokens = tokens[1:]
    while tokens and tokens[-1] in APPLICATION_DISCOVERY_COMPANY_SUFFIXES:
        tokens.pop()
    return " ".join(tokens)


def _discovery_title_tokens(value):
    tokens = set()
    for token in _discovery_normalize(value).split():
        if token.isdigit() or token in APPLICATION_DISCOVERY_TITLE_STOP_WORDS:
            continue
        tokens.add(APPLICATION_DISCOVERY_TITLE_ALIASES.get(token, token))
    return tokens


def _discovery_title_years(value):
    return set(re.findall(r"\b(?:19|20)\d{2}\b", _discovery_normalize(value)))


def _is_discovery_duplicate(existing_applications, company, title, source_url):
    company_key = _discovery_company_key(company)
    title_key = _discovery_normalize(title)
    title_tokens = _discovery_title_tokens(title)
    title_years = _discovery_title_years(title)

    for application in existing_applications:
        if application.url == source_url:
            return True
        existing_company = _discovery_company_key(application.company)
        existing_title = _discovery_normalize(application.title)
        if existing_company != company_key:
            continue
        if existing_title == title_key:
            return True
        # Confirmation emails often shorten a title slightly ("Software
        # Engineer Intern" vs "Software Engineering Internship"). Only use
        # containment for reasonably descriptive titles to avoid treating
        # every generic "Intern" application at a company as the same role.
        if min(len(existing_title), len(title_key)) >= 8 and (
            existing_title in title_key or title_key in existing_title
        ):
            return True
        # Assessment reminders commonly vary only by a year or by wording
        # such as "Internship Programme" vs "Intern Program". Token aliases
        # and a high overlap threshold catch those without collapsing two
        # genuinely different internships at the same employer.
        existing_tokens = _discovery_title_tokens(application.title)
        existing_years = _discovery_title_years(application.title)
        different_explicit_years = (
            existing_years and title_years and existing_years.isdisjoint(title_years)
        )
        if not different_explicit_years and min(len(existing_tokens), len(title_tokens)) >= 2:
            overlap = len(existing_tokens & title_tokens) / min(
                len(existing_tokens), len(title_tokens)
            )
            if overlap >= 0.8:
                return True
    return False


def _gmail_message_url(thread_id):
    # #all works for inbox, archived, and labelled messages alike.
    return f"https://mail.google.com/mail/u/0/#all/{thread_id}"


def _refresh_gmail_access_token(user):
    try:
        refresh_token = crypto.decrypt_token(user.gmail_refresh_token)
    except ValueError as e:
        current_app.logger.warning(f"Gmail token decrypt failed for user {user.email}: {e}")
        return None, (jsonify({"error": "Gmail connection is invalid -- please reconnect Gmail"}), 400)

    try:
        access_token = gmail_client.refresh_access_token(
            refresh_token,
            current_app.config.get("GOOGLE_CLIENT_ID"),
            current_app.config.get("GOOGLE_CLIENT_SECRET"),
        )
    except gmail_client.GmailScanError as e:
        current_app.logger.warning(f"Gmail token refresh failed for user {user.email}: {e}")
        return None, (jsonify({"error": "Gmail token refresh failed -- please reconnect Gmail"}), 502)
    return access_token, None


@api_routes.route("/discover-applications", methods=["POST"])
@token_required
def discover_applications():
    user = g.current_user
    if not user.gmail_refresh_token:
        return jsonify({"error": "Gmail not connected"}), 400

    data = request.get_json(silent=True) or {}
    try:
        lookback_days = int(data.get("lookback_days", 30))
    except (TypeError, ValueError):
        return jsonify({"error": "Choose a valid email search range"}), 400
    if lookback_days not in APPLICATION_DISCOVERY_LOOKBACK_DAYS:
        return jsonify({"error": "Search range must be 7, 30, 60, or 90 days"}), 400

    page_token = data.get("page_token")
    if page_token is not None and not isinstance(page_token, str):
        return jsonify({"error": "Invalid Gmail search page"}), 400

    access_token, error_response = _refresh_gmail_access_token(user)
    if error_response:
        return error_response

    since = date.today() - timedelta(days=lookback_days)
    query = f"after:{since.strftime('%Y/%m/%d')} {APPLICATION_DISCOVERY_GMAIL_QUERY}"
    try:
        # One bounded page keeps each request within Render's web-service
        # timeout. The dashboard follows next_page_token until every matching
        # email in the selected date range has been considered.
        message_ids, next_page_token = gmail_client.search_message_page(
            access_token,
            query,
            max_messages=APPLICATION_DISCOVERY_CANDIDATES_PER_PAGE,
            page_token=page_token,
        )
    except gmail_client.GmailScanError as e:
        current_app.logger.warning(f"Historical application search failed for user {user.email}: {e}")
        return jsonify({"error": "Gmail search failed"}), 502

    if not message_ids:
        return jsonify({
            "searched": 0,
            "added_applications": [],
            "skipped_duplicates": 0,
            "filtered_out": 0,
            "fetch_failures": 0,
            "next_page_token": next_page_token,
        })

    messages = []
    fetch_failures = 0
    for message_id in message_ids:
        try:
            metadata = gmail_client.get_message_details(access_token, message_id)
        except gmail_client.GmailScanError as e:
            current_app.logger.warning(f"Historical application email fetch failed for {message_id}: {e}")
            fetch_failures += 1
            continue

        received_date = datetime.fromtimestamp(
            metadata["internal_date"] / 1000, tz=timezone.utc
        ).date() if metadata["internal_date"] else since
        messages.append({
            **metadata,
            "received_date": received_date.isoformat(),
        })

    client = _get_genai_client()
    if client is None:
        return jsonify({"error": "Historical application sync unavailable"}), 503

    discovered = []
    for start in range(0, len(messages), APPLICATION_DISCOVERY_BATCH_SIZE):
        batch = messages[start:start + APPLICATION_DISCOVERY_BATCH_SIZE]
        allowed_ids = {message["id"] for message in batch}
        try:
            result = client.models.generate_content(
                model=GEMINI_MODEL,
                contents=_build_application_discovery_prompt(batch),
                config=genai_types.GenerateContentConfig(
                    http_options=genai_types.HttpOptions(timeout=GEMINI_TIMEOUT_MS),
                    response_mime_type="application/json",
                ),
            )
            parsed = _parse_application_discovery_response(result.text or "", allowed_ids)
        except Exception as e:
            current_app.logger.warning(f"Historical application discovery call failed: {e}")
            return jsonify({"error": "Couldn't analyse those emails -- try again"}), 503
        if parsed is None:
            return jsonify({"error": "Couldn't analyse those emails -- try again"}), 503
        discovered.extend(parsed)

    discovered_message_ids = {item["message_id"] for item in discovered}
    filtered_out = sum(
        1 for message in messages if message["id"] not in discovered_message_ids
    )

    metadata_by_id = {message["id"]: message for message in messages}
    existing_applications = Application.query.filter_by(user_id=user.email).all()
    added = []
    skipped_duplicates = 0

    for item in discovered:
        metadata = metadata_by_id[item["message_id"]]
        source_url = _gmail_message_url(
            metadata.get("thread_id") or item["message_id"]
        )
        if _is_discovery_duplicate(
            existing_applications, item["company"], item["title"], source_url
        ):
            skipped_duplicates += 1
            continue

        parsed_date = _parse_date(item["applied_date"])
        fallback_date = _parse_date(metadata["received_date"]) or since
        if parsed_date is None or parsed_date < since or parsed_date > date.today():
            parsed_date = fallback_date

        application = Application(
            user_id=user.email,
            title=item["title"],
            company=item["company"],
            url=source_url,
            status=ApplicationStatus.APPLIED,
            applied_date=parsed_date,
            notes="Discovered from a Gmail application or recruitment-progression email.",
        )
        application.set_flags([])
        db.session.add(application)
        db.session.flush()
        existing_applications.append(application)
        added.append(application)

    db.session.commit()
    return jsonify({
        "searched": len(messages),
        "added_applications": [application.to_dict() for application in added],
        "skipped_duplicates": skipped_duplicates,
        "filtered_out": filtered_out,
        "fetch_failures": fetch_failures,
        "next_page_token": next_page_token,
    })


@api_routes.route("/scan-emails", methods=["POST"])
@token_required
def scan_emails():
    user = g.current_user

    if not user.gmail_refresh_token:
        return jsonify({"error": "Gmail not connected"}), 400

    applications = Application.query.filter_by(user_id=user.email).all()

    if user.last_email_scan_at:
        since = user.last_email_scan_at.date()
    else:
        applied_dates = [a.applied_date for a in applications if a.applied_date]
        if not applied_dates:
            # Nothing tracked yet to match against -- there's no meaningful
            # window to search, and no way for anything to match anyway.
            return jsonify({"scanned": 0, "updated_applications": []})
        since = min(applied_dates)

    try:
        refresh_token = crypto.decrypt_token(user.gmail_refresh_token)
    except ValueError as e:
        current_app.logger.warning(f"Gmail token decrypt failed for user {user.email}: {e}")
        return jsonify({"error": "Gmail connection is invalid -- please reconnect Gmail"}), 400

    client_id = current_app.config.get("GOOGLE_CLIENT_ID")
    client_secret = current_app.config.get("GOOGLE_CLIENT_SECRET")

    try:
        access_token = gmail_client.refresh_access_token(refresh_token, client_id, client_secret)
    except gmail_client.GmailScanError as e:
        current_app.logger.warning(f"Gmail token refresh failed for user {user.email}: {e}")
        return jsonify({"error": "Gmail token refresh failed -- please reconnect Gmail"}), 502

    query = f"after:{since.strftime('%Y/%m/%d')}"

    try:
        message_ids = gmail_client.search_message_ids(access_token, query)
    except gmail_client.GmailScanError as e:
        current_app.logger.warning(f"Gmail search failed for user {user.email}: {e}")
        return jsonify({"error": "Gmail search failed"}), 502

    if not message_ids:
        user.last_email_scan_at = datetime.utcnow()
        db.session.commit()
        return jsonify({"scanned": 0, "updated_applications": []})

    already_processed_ids = {
        row.gmail_message_id
        for row in ProcessedEmail.query.filter(
            ProcessedEmail.user_id == user.id,
            ProcessedEmail.gmail_message_id.in_(message_ids),
        ).all()
    }
    candidate_ids = [mid for mid in message_ids if mid not in already_processed_ids]

    genai_client_instance = _get_genai_client()

    # Phase 1: cheap metadata-only fetch + match decision for every
    # candidate. Unmatched candidates are fully resolved here (no body fetch,
    # no Gemini call needed) and recorded immediately.
    matched_candidates = []
    scanned_count = 0

    for message_id in candidate_ids:
        try:
            metadata = gmail_client.get_message_metadata(access_token, message_id)
        except gmail_client.GmailScanError as e:
            current_app.logger.warning(f"Gmail metadata fetch failed for message {message_id}: {e}")
            continue  # left unprocessed -- retried on the next scan

        scanned_count += 1
        matched_app = email_matching.find_best_match(applications, metadata["subject"], metadata["sender"])

        if matched_app is None:
            db.session.add(ProcessedEmail(user_id=user.id, gmail_message_id=message_id, application_id=None))
            db.session.commit()
            continue

        matched_candidates.append((matched_app, metadata))

    # Phase 2: process actual matches in chronological order (oldest first),
    # so if two emails both match the same application, the one that's
    # genuinely newer is the one left standing -- regardless of what order
    # Gmail's search API or this loop happened to encounter them in.
    matched_candidates.sort(key=lambda pair: pair[1]["internal_date"])

    # Keyed by application id so an application overwritten twice in the same
    # scan (two matching emails, see the chronological sort above) is only
    # reported once in the response, reflecting its final state.
    updated_applications_by_id = {}

    for matched_app, metadata in matched_candidates:
        message_id = metadata["id"]

        try:
            body = gmail_client.get_message_body(access_token, message_id)
        except gmail_client.GmailScanError as e:
            current_app.logger.warning(f"Gmail body fetch failed for message {message_id}: {e}")
            continue  # left unprocessed -- retried on the next scan

        if genai_client_instance is None:
            current_app.logger.warning("Gmail scan classification skipped: GEMINI_API_KEY not configured")
            continue  # left unprocessed -- retried once Gemini is configured

        try:
            result = genai_client_instance.models.generate_content(
                model=GEMINI_MODEL,
                contents=_build_classification_prompt(metadata["subject"], body),
                config=genai_types.GenerateContentConfig(
                    http_options=genai_types.HttpOptions(timeout=GEMINI_TIMEOUT_MS),
                    response_mime_type="application/json",
                ),
            )
            raw_text = (result.text or "").strip()
        except Exception as e:
            current_app.logger.warning(f"Gemini classification failed for message {message_id}: {e}")
            continue  # left unprocessed -- retried on the next scan

        classification = _parse_classification_response(raw_text)
        if classification is None:
            current_app.logger.warning(
                f"Gemini classification for message {message_id} wasn't a valid label -- treating as Not Relevant"
            )
            classification = "Not Relevant"

        db.session.add(ProcessedEmail(user_id=user.id, gmail_message_id=message_id, application_id=matched_app.id))

        if classification in SUGGESTION_CATEGORIES:
            matched_app.ai_suggested_status = classification
            matched_app.ai_suggestion_source_email_id = message_id
            matched_app.ai_suggestion_seen = False
            matched_app.ai_suggestion_created_at = datetime.utcnow()
            updated_applications_by_id[matched_app.id] = matched_app

        db.session.commit()

    # Only advance the watermark once the scan actually completed --
    # anything skipped above (transient failures) stayed out of
    # ProcessedEmail, so it's naturally retried on the next scan rather than
    # silently lost in a gap.
    user.last_email_scan_at = datetime.utcnow()
    db.session.commit()

    return jsonify({
        "scanned": scanned_count,
        "updated_applications": [a.to_dict() for a in updated_applications_by_id.values()],
    })
